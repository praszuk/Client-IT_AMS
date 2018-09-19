import logging

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from Main import CONFIG

TEMPLATE_FILE = CONFIG.TEMPLATE_FILE_PATH

DATE_START_TAG = '<date_start>'
DATE_END_TAG = '<date_end>'

CLIENT_ADDRESS_TAG = '<client_info>'
CLIENT_NAME = '<client>'

MIDDLEMAN_NAME = '<middleman>'


class Generator:
    """
        This class will generate protocols for loans based on a rendered template with CONSTANT fields.
        Class is not flexible, it is adjusted to one type of protocol which basing on structure:
            - few fields
            - 2 tables where second contains 1 cell with table inside (nested)
    """

    def __init__(self, template=TEMPLATE_FILE):
        """
        :param template: path to template (docx file)
        :type template: str
        """

        self.__template = template
        self.__document = None

    def generate_loan_protocol(self, checkout_model, hardware):

        """
        Generate loan protocol using static template from file TEMPLATE_FILE

        :param checkout_model: model contains middleman name, client data and dates
        :type checkout_model: CheckoutModel

        :param hardware: hardware products list. List of Assets.
        :type hardware: list of AssetModel
        """

        middleman_name = checkout_model.middle_man_name
        client_name = checkout_model.client_name
        client_address = checkout_model.client_address
        date_start = checkout_model.checkout_date
        date_end = checkout_model.check_in_date

        # Create doc
        self.__document = Document(self.__template)

        # Add hardware data to table[0]
        for asset in hardware:
            cells = self.__document.tables[0].add_row().cells

            # Remove default paragraph in cell and create new with appropriate alignment
            for cell in cells:
                self.delete_paragraph(cell.paragraphs[-1])
                _p = cell.add_paragraph('')
                _p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Just setting text to cell (It's important to set it to [0] - first paragraph. With alignment def. above)
            cells[0].paragraphs[0].text = asset.name
            cells[1].paragraphs[0].text = asset.notes
            cells[2].paragraphs[0].text = asset.serial_number

        # Add signatures (nested table)
        for row in self.__document.tables[1].rows[0].cells[0].tables[0].rows:
            for cell in row.cells:
                if CLIENT_NAME in cell.text:
                    cell.text = cell.text.replace(CLIENT_NAME, client_name)

                elif MIDDLEMAN_NAME in cell.text:
                    cell.text = cell.text.replace(MIDDLEMAN_NAME, middleman_name)
                    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Insert dates and client address
        for idx, p in enumerate(self.__document.paragraphs):
            if DATE_START_TAG in p.text:
                curr_text = p.text
                p.clear()

                start_index = curr_text.find(DATE_START_TAG)

                p.add_run(curr_text[:start_index])
                p.add_run(date_start).bold = True
                p.add_run(curr_text[start_index + len(DATE_START_TAG):])

            elif DATE_END_TAG in p.text:
                curr_text = p.text
                p.clear()

                start_index = curr_text.find(DATE_END_TAG)
                p.add_run(curr_text[:start_index - 1]).underline = True

                run = p.add_run(' ' + date_end)
                run.bold = True
                run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)  # RED

            elif CLIENT_ADDRESS_TAG in p.text:
                p.text = client_address

    def save_to_file(self, filename):
        self.__document.save(filename)
        logging.info('Saved document at {}'.format(filename))

    @staticmethod
    def delete_paragraph(paragraph):
        """
        Remove paragraph in 'safety way'
        https://github.com/python-openxml/python-docx/issues/33#issuecomment-84706929
        """
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
