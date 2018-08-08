from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor

from Main import CONFIG

TEMPLATE_FILE = CONFIG.TEMPLATE_FILE_PATH

DATE_START_TAG = '<date_start>'
DATE_END_TAG = '<date_end>'

CLIENT_ADDRESS_TAG = '<client_info>'
CLIENT_NAME = '<client>'

MIDDLEMAN_NAME = '<middleman>'


class Generator:
    """
        This class will be generate protocols for loans based on rendered template with CONSTANT fields.
        Class is not flexible, it is adjusted to one type of protocol which basing on structure:
            - few fields
            - 2 tables where second contains 1 cell with table inside
        :param template: path to template docx file
        :type template: str
    """

    def __init__(self, template=TEMPLATE_FILE):
        self.__document = Document(template)

    def generate_loan_protocol(self, middleman_name, client_name, client_address, date_start, date_end, hardware):
        """
        Generate loan protocol using static template from file '.template.docx'

        :param middleman_name: first name and last name
        :type middleman_name: str

        :param client_name: first name and last name
        :type client_name: str

        :param client_address: home/business/company address
        :type client_address: str

        :param date_start: begin of checkout
        :type date_start: str

        :param date_end: expected day of check-in
        :type date_end: str

        :param hardware: hardware products list. Each element of list contains [name, description, serial_number]
        :type hardware: list of (list [str, str, str])
        """

        # Add hardware data to table[0]
        for name, desc, serial_number in hardware:
            cells = self.__document.tables[0].add_row().cells
            cells[0].text = name
            cells[1].text = desc
            cells[2].text = serial_number

        # Add signatures (nested table)
        for row in self.__document.tables[1].rows[0].cells[0].tables[0].rows:
            for cell in row.cells:
                if CLIENT_NAME in cell.text:
                    cell.text = cell.text.replace(CLIENT_NAME, client_name)

                elif MIDDLEMAN_NAME in cell.text:
                    cell.text = cell.text.replace(MIDDLEMAN_NAME, middleman_name)
                    cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Insert dates and client address
        for idx, p in enumerate(self.__document.paragraphs):
            if DATE_START_TAG in p.text:
                curr_text = p.text
                p.clear()

                start_index = curr_text.find(DATE_START_TAG)

                p.add_run(curr_text[:start_index])
                p.add_run(date_start).bold = True
                p.add_run(curr_text[start_index + len(DATE_START_TAG):])

            elif DATE_END_TAG in p.text:
                curr_text = p.text
                p.clear()

                start_index = curr_text.find(DATE_END_TAG)
                p.add_run(curr_text[:start_index - 1]).underline = True

                run = p.add_run(' ' + date_end)
                run.bold = True
                run.font.color.rgb = RGBColor(0xff, 0x00, 0x00)  # RED

            elif CLIENT_ADDRESS_TAG in p.text:
                p.text = client_address

    def save_to_file(self, filename):
        self.__document.save(filename)
